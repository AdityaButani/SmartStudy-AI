import streamlit as st
import os
from dotenv import load_dotenv
from groq import Groq
import datetime
import docx
from docx.shared import Inches
from io import BytesIO
import base64

# Page configuration
st.set_page_config(
    page_title="AI Study Plan Generator",
    page_icon="📚",
    layout="wide"
)

# Custom CSS
st.markdown("""
    <style>
    /* Main container styling */
    .stApp {
        background-color: #f8f9fa;
    }
    
    /* Card styling */
    .styledcard {
        background-color: white;
        padding: 2rem;
        border-radius: 1rem;
        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);
        margin-bottom: 1rem;
    }
    
    /* Header styling */
    .main-header {
        color: #1e3a8a;
        font-size: 2.5rem;
        font-weight: bold;
        margin-bottom: 1.5rem;
        text-align: center;
        padding: 1rem 0;
    }
    
    /* Section headers */
    .section-header {
        color: #1e40af;
        font-size: 1.5rem;
        font-weight: 600;
        margin-bottom: 1rem;
    }
    
    /* Input field labels */
    .input-label {
        font-weight: 500;
        color: #374151;
        margin-bottom: 0.5rem;
    }
    
    /* Generated plan styling */
    .study-plan {
        background-color: #ffffff;
        padding: 2rem;
        border-radius: 0.5rem;
        border-left: 4px solid #2563eb;
        margin: 1rem 0;
    }
    
    /* Download button styling */
    .download-button {
        background-color: #2563eb;
        color: white;
        padding: 0.75rem 1.5rem;
        border-radius: 0.5rem;
        text-decoration: none;
        display: inline-block;
        margin-top: 1rem;
        font-weight: 500;
    }
    
    .download-button:hover {
        background-color: #1d4ed8;
    }
    
    /* Progress bar styling */
    .stProgress > div > div {
        background-color: #2563eb;
    }
    
    /* Tip box styling */
    .tip-box {
        background-color: #dbeafe;
        border-left: 4px solid #2563eb;
        padding: 1rem;
        border-radius: 0.5rem;
        margin: 1rem 0;
    }
    </style>
""", unsafe_allow_html=True)

# Load environment variables
load_dotenv()
groq_api_key = os.environ.get("GROQ_API_KEY")

if not groq_api_key:
    st.error("❌ GROQ_API_KEY environment variable not set. Please set it.")
    st.stop()

client = Groq(api_key=groq_api_key)

# Header
st.markdown('<h1 class="main-header">📚 PrepMaster - AI Study Plan Generator</h1>', unsafe_allow_html=True)

# Introduction
st.markdown("""
    <div class="tip-box">
        <h4>🌟 Create Your Personalized Learning Journey</h4>
        <p>Input your learning goals and preferences below, and our AI will generate a customized study plan tailored just for you.</p>
    </div>
""", unsafe_allow_html=True)

# Create two columns for the form
col1, col2 = st.columns(2)

with col1:
    st.markdown('<div class="styledcard">', unsafe_allow_html=True)
    st.markdown('<h3 class="section-header">📝 Basic Information</h3>', unsafe_allow_html=True)
    subject = st.text_input("What subject do you want to study?", 
                           placeholder="e.g., Python Programming, Data Science, Machine Learning")
    
    current_knowledge = st.text_area("What is your current knowledge level?",
                                   placeholder="e.g., Beginner with basic programming concepts",
                                   height=120)
    
    learning_goals = st.text_area("What are your learning goals?",
                                 placeholder="e.g., Build web applications using Python and Django",
                                 height=120)
    st.markdown('</div>', unsafe_allow_html=True)

with col2:
    st.markdown('<div class="styledcard">', unsafe_allow_html=True)
    st.markdown('<h3 class="section-header">⚙️ Study Preferences</h3>', unsafe_allow_html=True)
    
    duration_weeks = st.slider("Study Duration (weeks):", 
                             min_value=1, max_value=52, value=8,
                             help="Select the number of weeks you want to study")
    
    time_per_week = st.slider("Hours per Week:", 
                             min_value=1, max_value=40, value=10,
                             help="Select how many hours you can dedicate per week")
    
    resource_preference = st.selectbox("Preferred Learning Resources:",
                                     ["Any", "Textbooks", "Online Courses", "Video Tutorials", 
                                      "Research Papers", "Interactive Projects"],
                                     help="Select your preferred type of learning materials")
    
    style_preference = st.selectbox("Learning Style:",
                                  ["Any", "Structured", "Flexible", "Practical", "Theoretical"],
                                  help="Select your preferred approach to learning")
    
    difficulty_level = st.select_slider("Preferred Content Difficulty:",
                                      options=["Beginner", "Intermediate", "Advanced"],
                                      value="Intermediate",
                                      help="Select the difficulty level of the content")
    st.markdown('</div>', unsafe_allow_html=True)

# Generate button
st.markdown('<div class="styledcard">', unsafe_allow_html=True)
generate_button = st.button("🚀 Generate My Study Plan", use_container_width=True)

if generate_button:
    if not subject or not current_knowledge or not learning_goals:
        st.warning("🚨 Please fill in all the required fields.")
    else:
        try:
            with st.spinner("🎨 Creating your personalized study plan..."):
                # Construct the prompt
                prompt = f"""
                Subject: {subject}
                Current Knowledge: {current_knowledge}
                Learning Goals: {learning_goals}
                Duration: {duration_weeks} weeks, {time_per_week} hours per week
                Preferred Resources: {resource_preference}
                Learning Style: {style_preference}
                Difficulty Level: {difficulty_level}

                Please create a detailed, week-by-week study plan that includes:
                1. Weekly learning objectives
                2. Specific topics to cover
                3. Recommended learning resources (with links where applicable)
                4. Practice exercises or projects
                5. Progress tracking milestones
                
                Format the plan with clear headers, bullet points, and a logical progression.
                Include time estimates for each activity.
                """

                # Call the Groq API
                chat_completion = client.chat.completions.create(
                    messages=[{"role": "user", "content": prompt}],
                    model="llama-3.3-70b-versatile",
                    max_tokens=2048,
                )

                study_plan = chat_completion.choices[0].message.content

                # Display the generated plan
                st.markdown('<div class="study-plan">', unsafe_allow_html=True)
                st.markdown('<h3 class="section-header">📋 Your Personalized Study Plan</h3>', 
                          unsafe_allow_html=True)
                st.write(study_plan)
                st.markdown('</div>', unsafe_allow_html=True)

                # Create Word document
                doc = docx.Document()
                doc.add_heading(f"Personalized Study Plan: {subject}", 0)
                doc.add_heading("Study Details", level=1)
                doc.add_paragraph(f"Generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")
                doc.add_paragraph(f"Subject: {subject}")
                doc.add_paragraph(f"Duration: {duration_weeks} weeks, {time_per_week} hours/week")
                doc.add_paragraph(f"Learning Style: {style_preference}")
                doc.add_paragraph(f"Difficulty Level: {difficulty_level}")
                doc.add_paragraph(f"Resource Preference: {resource_preference}")
                
                doc.add_heading("Learning Goals", level=1)
                doc.add_paragraph(learning_goals)
                
                doc.add_heading("Current Knowledge Level", level=1)
                doc.add_paragraph(current_knowledge)
                
                doc.add_heading("Study Plan", level=1)
                doc.add_paragraph(study_plan)

                # Save and create download button
                doc_stream = BytesIO()
                doc.save(doc_stream)
                doc_stream.seek(0)
                
                b64 = base64.b64encode(doc_stream.read()).decode()
                href = f'''
                    <a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" 
                    download="study_plan.docx" class="download-button">
                    📥 Download Study Plan (Word Document)
                    </a>
                '''
                st.markdown(href, unsafe_allow_html=True)

        except Exception as e:
            st.error(f"❌ An error occurred: {str(e)}")

st.markdown('</div>', unsafe_allow_html=True)

# Footer with tips
st.markdown("""
    <div class="tip-box">
        <h4>💡 Tips for Better Results:</h4>
        <ul>
            <li>Be specific about your current knowledge level and goals</li>
            <li>Consider breaking down long-term goals into smaller milestones</li>
            <li>Be realistic about the time you can dedicate to studying</li>
            <li>Download and save your study plan for offline reference</li>
        </ul>
    </div>
""", unsafe_allow_html=True)