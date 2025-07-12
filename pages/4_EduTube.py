import streamlit as st
import os
import time
import google.generativeai as genai
from dotenv import load_dotenv
import tiktoken
import concurrent.futures
from youtube_transcript_api import YouTubeTranscriptApi
from docx import Document
import json
import pandas as pd
from datetime import datetime
import re
from streamlit_option_menu import option_menu
import random

# Set page config at the very top
st.set_page_config(
    page_title="EduTube AI Assistant",
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Define CHUNK_SIZE based on the model being used
MODEL_NAME = "gemini-1.5-pro"  # Default model
CHUNK_SIZE = 200000 if MODEL_NAME == "gemini-1.5-pro" else 100000

# Load environment variables
load_dotenv()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

if not GEMINI_API_KEY:
    st.error("GEMINI_API_KEY environment variable is missing. Please set it in your .env file.")
    st.stop()

# Initialize Gemini client
genai.configure(api_key=GEMINI_API_KEY)

# Set up Gemini model configuration
model = genai.GenerativeModel(MODEL_NAME)

# Enhanced Custom CSS for better styling
st.markdown("""
    <style>
    /* Global Styles */
    .stApp {
        background: linear-gradient(135deg, #f5f7fa 0%, #e4e9f2 100%);
    }

    /* Header Styles */
    .quiz-header {
        text-align: center;
        margin: 2rem 0;
        padding: 2rem;
        background: linear-gradient(45deg, #2b5876 0%, #4e4376 100%);
        border-radius: 15px;
        color: white;
        box-shadow: 0 4px 15px rgba(0,0,0,0.1);
    }

    .quiz-header h1 {
        font-size: 2.5rem;
        font-weight: 700;
        text-shadow: 2px 2px 4px rgba(0,0,0,0.2);
    }

    /* Question Card Styles */
    .question-card {
        background: white;
        padding: 25px;
        border-radius: 15px;
        margin: 20px 0;
        border: none;
        box-shadow: 0 4px 15px rgba(0,0,0,0.05);
        transition: transform 0.2s ease;
    }

    .question-card:hover {
        transform: translateY(-2px);
    }

    .question-card h4 {
        color: #2b5876;
        font-size: 1.2rem;
        margin-bottom: 15px;
    }

    /* Answer Styles */
    .correct-answer {
        color: #10b981;
        font-weight: 600;
        padding: 10px;
        border-radius: 8px;
        background: #ecfdf5;
    }

    .incorrect-answer {
        color: #ef4444;
        font-weight: 600;
        padding: 10px;
        border-radius: 8px;
        background: #fef2f2;
    }

    /* Result Card Styles */
    .result-card {
        background: white;
        padding: 30px;
        border-radius: 15px;
        margin: 20px 0;
        box-shadow: 0 4px 15px rgba(0,0,0,0.1);
    }

    /* Custom Button Styles */
    .stButton button {
        background: linear-gradient(45deg, #2b5876 0%, #4e4376 100%);
        color: white;
        border: none;
        padding: 0.5rem 2rem;
        border-radius: 8px;
        font-weight: 600;
        transition: all 0.3s ease;
    }

    .stButton button:hover {
        transform: translateY(-2px);
        box-shadow: 0 4px 15px rgba(0,0,0,0.2);
    }

    /* Input Field Styles */
    .stTextInput input, .stNumberInput input {
        border-radius: 8px;
        border: 2px solid #e4e9f2;
        padding: 10px;
        transition: all 0.3s ease;
    }

    .stTextInput input:focus, .stNumberInput input:focus {
        border-color: #2b5876;
        box-shadow: 0 0 0 2px rgba(43,88,118,0.2);
    }

    /* Radio Button Styles */
    .stRadio label {
        padding: 10px;
        border-radius: 8px;
        transition: all 0.2s ease;
    }

    .stRadio label:hover {
        background: #f8fafc;
    }

    /* Progress Bar Styles */
    .stProgress > div > div {
        background-color: #2b5876;
    }

    /* Score Display Styles */
    .final-score {
        font-size: 2rem;
        font-weight: 700;
        color: #2b5876;
        text-align: center;
        padding: 20px;
        background: white;
        border-radius: 15px;
        box-shadow: 0 4px 15px rgba(0,0,0,0.1);
    }

    /* Divider Styles */
    hr {
        border: none;
        height: 2px;
        background: linear-gradient(90deg, transparent, #e4e9f2, transparent);
        margin: 25px 0;
    }
    </style>
    """, unsafe_allow_html=True)

# Initialize all session state variables
if "quiz_data" not in st.session_state:
    st.session_state.quiz_data = []
if "user_answers" not in st.session_state:
    st.session_state.user_answers = {}
if "quiz_generated" not in st.session_state:
    st.session_state.quiz_generated = False
if "quiz_submitted" not in st.session_state:
    st.session_state.quiz_submitted = False
if 'history' not in st.session_state:
    st.session_state.history = []
if 'current_summary' not in st.session_state:
    st.session_state.current_summary = ""
if 'current_notes' not in st.session_state:
    st.session_state.current_notes = ""
if 'current_quiz' not in st.session_state:
    st.session_state.current_quiz = ""
if 'video_title' not in st.session_state:
    st.session_state.video_title = ""

def extract_video_id(url):
    """Extract YouTube video ID from various URL formats"""
    patterns = [
        r"(?:youtube\.com\/watch\?v=|youtu\.be\/)([\w-]+)",
        r"(?:youtube\.com\/embed\/)([\w-]+)",
        r"(?:youtube\.com\/v\/)([\w-]+)",
    ]
    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            return match.group(1)
    return None

def get_video_title(video_id):
    """Get the title of a YouTube video"""
    try:
        import requests
        from bs4 import BeautifulSoup
        url = f"https://www.youtube.com/watch?v={video_id}"
        response = requests.get(url)
        soup = BeautifulSoup(response.text, 'html.parser')
        title = soup.find('title').text.replace(' - YouTube', '')
        return title
    except Exception:
        return f"Video {video_id}"

def get_transcript(video_id):
    """Get transcript from YouTube video"""
    try:
        transcript_list = YouTubeTranscriptApi.get_transcript(video_id, languages=["en"])
        return " ".join([t["text"] for t in transcript_list])
    except Exception as e:
        st.error(f"Error fetching transcript: {e}")
        return None

def chunk_text(text, chunk_size):
    """Split text into chunks of specified token size"""
    encoding = tiktoken.encoding_for_model("gpt-3.5-turbo")  # Using compatible tokenizer
    tokens = encoding.encode(text)
    chunks = [tokens[i:i + chunk_size] for i in range(0, len(tokens), chunk_size)]
    return [encoding.decode(chunk) for chunk in chunks]

def process_with_gemini(transcript, detail_level, output_type, model_name):
    """Process transcript directly with Gemini, with fewer chunks due to higher token limits"""
    try:
        if len(transcript) > 30000:  # If still too large, break into chunks
            transcript_chunks = chunk_text(transcript, CHUNK_SIZE)
            results = process_chunks_with_gemini(transcript_chunks, detail_level, output_type, model_name)
            return results
        else:
            # Process directly without chunking
            result = generate_content_with_gemini(transcript, detail_level, output_type, model_name)
            return result
    except Exception as e:
        st.error(f"Error processing with Gemini: {e}")
        return "Error processing content. Please try again."

def process_chunks_with_gemini(transcript_chunks, detail_level, output_type, model_name):
    """Process multiple chunks with Gemini and merge results"""
    results = []
    total_chunks = len(transcript_chunks)

    progress_bar = st.progress(0)
    status_text = st.empty()

    def process_chunk_with_retry(chunk, detail_level, output_type, model_name, retries=3):
        for attempt in range(retries):
            try:
                return generate_content_with_gemini(chunk, detail_level, output_type, model_name)
            except Exception as e:
                st.error(f"Error: {e}. Retrying in {2 ** attempt} seconds...")
                time.sleep(2 ** attempt + random.uniform(0, 1))  # Exponential backoff
        raise RuntimeError("Failed to process chunk after multiple attempts.")

    with concurrent.futures.ThreadPoolExecutor() as executor:
        future_to_chunk = {executor.submit(process_chunk_with_retry, chunk, detail_level, output_type, model_name): chunk for chunk in transcript_chunks}
        for i, future in enumerate(concurrent.futures.as_completed(future_to_chunk)):
            try:
                result = future.result()
                results.append(result)
                st.write(f"Chunk {i+1} processed successfully.")
            except Exception as exc:
                st.error(f"Chunk {i+1} generated an exception: {exc}")
            progress_bar.progress((i + 1) / total_chunks)
            status_text.text(f"Processing chunk {i+1}/{total_chunks}...")

    status_text.text("Merging results...")

    # For summaries, notes, and quizzes, we need to merge intelligently
    merged_result = merge_with_gemini(results, output_type, model_name)

    progress_bar.progress(1.0)
    status_text.text("Processing complete!")

    return merged_result

def generate_content_with_gemini(text, detail_level, output_type, model_name):
    """Generate content using Gemini API based on output type"""

    if output_type == "summary":
        prompt = f"""
        Please provide a comprehensive analysis and summary of this educational content using the following structure:

        1. SUBJECT/DOMAIN:
        - Identify the main subject area or field this content belongs to.

        2. MAIN TOPICS COVERED:
        - List key topics or concepts discussed
        - Include important terms or definitions mentioned

        3. DETAILED SUMMARY:
        - Provide a structured summary of the main content
        - Highlight key points, arguments, and examples

        4. KEY TAKEAWAYS:
        - List the most important lessons or insights

        Detail Level: {detail_level.capitalize()}

        Text to analyze: {text}
        """
    elif output_type == "notes":
        prompt = f"""
        Create organized study notes from this educational content using the following structure:

        1. TITLE:
        - Create a descriptive title for these notes

        2. KEY CONCEPTS:
        - Bullet points of the most important concepts
        - Include definitions of important terms

        3. DETAILED NOTES:
        - Organized with headings and subheadings
        - Include bullet points for important details
        - Add any formulas, theories, or frameworks mentioned
        - Provide in-depth explanations and examples for each concept

        4. EXAMPLES:
        - Note any examples mentioned that illustrate key concepts

        5. QUESTIONS TO CONSIDER:
        - Create 2-3 reflective questions based on the content

        Detail Level: Deep Detail

        Text to analyze: {text}
        """
    elif output_type == "quiz":
        prompt = f"""
        Create a quiz based on this educational content with the following components. Format your response EXACTLY according to this structure:

        # MULTIPLE CHOICE QUESTIONS
        1. [Question text]
        A) [Option A]
        B) [Option B]
        C) [Option C]
        D) [Option D]
        Correct: [A/B/C/D]

        2. [Question text]
        A) [Option A]
        B) [Option B]
        C) [Option C]
        D) [Option D]
        Correct: [A/B/C/D]

        [Continue for 5 questions total]

        # TRUE/FALSE QUESTIONS
        1. [Statement]
        Correct: [True/False]

        2. [Statement]
        Correct: [True/False]

        [Continue for 3 questions total]

        # SHORT ANSWER QUESTIONS
        1. [Question text]
        Answer: [Brief model answer]

        2. [Question text]
        Answer: [Brief model answer]

        Make sure all questions are challenging but fair, based on important concepts from the content.
        Detail Level: {detail_level.capitalize()}

        Text to analyze: {text}
        """

    try:
        model = genai.GenerativeModel(model_name)
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        st.error(f"Gemini API Error: {e}")
        return f"Error generating {output_type}: {str(e)}"

def merge_with_gemini(content_parts, output_type, model_name):
    """Merge multiple content parts using Gemini"""

    # Join content parts with separators
    combined = "\n\n===CHUNK SEPARATOR===\n\n".join(content_parts)

    if output_type == "quiz":
        prompt = """
        I have quiz content from multiple chunks that needs to be merged into one coherent quiz.
        Please organize this content into a single quiz with:

        1. 5 Multiple choice questions (with options A, B, C, D and marked correct answers)
        2. 3 True/False questions (with marked correct answers)
        3. 2 Short answer questions (with model answers)

        Remove any duplicates or very similar questions. Ensure the final quiz follows exactly this format:

        # MULTIPLE CHOICE QUESTIONS
        1. [Question text]
        A) [Option A]
        B) [Option B]
        C) [Option C]
        D) [Option D]
        Correct: [A/B/C/D]

        [Continue for exactly 5 questions total]

        # TRUE/FALSE QUESTIONS
        1. [Statement]
        Correct: [True/False]

        [Continue for exactly 3 questions total]

        # SHORT ANSWER QUESTIONS
        1. [Question text]
        Answer: [Brief model answer]

        [Continue for exactly 2 questions total]

        Content to merge:
        """
    else:
        prompt = """
        I have content from multiple chunks of the same video that need to be merged intelligently.
        Please organize this content into a cohesive, non-repetitive document that maintains the
        structure but removes duplications. Merge similar sections, remove redundancies, and
        create a clean, unified document.

        Content to merge:
        """

    prompt += combined

    try:
        model = genai.GenerativeModel(model_name)
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        st.error(f"Error merging content: {e}")
        return "\n\n".join(content_parts)  # Fallback to simple concatenation

def save_to_word(content, file_name="document.docx", title="Document"):
    """Save content to a Word document"""
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(content)
    file_path = file_name
    doc.save(file_path)
    return file_path

def add_to_history(url, video_title, output_type):
    """Add processed video to history"""
    # Add to session state history
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")
    st.session_state.history.append({
        "timestamp": timestamp,
        "url": url,
        "title": video_title,
        "type": output_type
    })

    # Save to file for persistence
    try:
        with open("history.json", "w") as f:
            json.dump(st.session_state.history, f)
    except:
        pass

def load_history():
    """Load history from file"""
    try:
        with open("history.json", "r") as f:
            st.session_state.history = json.load(f)
    except:
        st.session_state.history = []

# UI Components
def render_sidebar():
    with st.sidebar:
        st.image("https://via.placeholder.com/150x150.png?text=EduTube", width=150)
        st.title("EduTube AI")

        selected = option_menu(
            "Navigation",
            ["Summarizer", "History", "Settings"],
            icons=["book", "clock-history", "gear"],
            menu_icon="cast",
            default_index=0,
        )

        st.markdown("---")
        st.markdown("### About")
        st.markdown("EduTube AI helps you extract knowledge from educational YouTube videos with AI-powered summaries, notes, and quizzes.")

        # LLM Selection in Sidebar
        st.session_state.model_name = st.selectbox(
            "Choose LLM",
            options=["gemini-1.5-pro", "gemini-1.5-flash"],
            index=0
        )

        return selected

def render_header():
    col1, col2 = st.columns([1, 4])
    with col1:
        st.image("https://via.placeholder.com/80x80.png?text=📚", width=80)
    with col2:
        st.title("EduTube AI Assistant")
        st.markdown("Transform YouTube lectures into structured learning resources")

def render_main_app():
    # Input Fields in a card-like container
    with st.container():
        st.markdown("### 🎬 Enter Video Details")
        url = st.text_input("YouTube Video URL:", placeholder="https://www.youtube.com/watch?v=...")

        col1, col2 = st.columns(2)
        with col1:
            detail_level = st.select_slider(
                "Detail Level",
                options=["Concise", "Detailed", "Deep Detail"],
                value="Deep Detail"
            )
        with col2:
            output_type = st.radio(
                "Output Type",
                options=["Summary", "Study Notes", "Quiz"],
                horizontal=True
            )

        output_map = {
            "Summary": "summary",
            "Study Notes": "notes",
            "Quiz": "quiz"
        }

        process_button = st.button("Process Video", type="primary", use_container_width=True)

    # Process the video if button is clicked
    if url and process_button:
        video_id = extract_video_id(url)
        if video_id:
            # Get video title
            video_title = get_video_title(video_id)
            st.session_state.video_title = video_title

            # Display video
            st.video(url)

            with st.spinner("Fetching transcript..."):
                transcript = get_transcript(video_id)
                if transcript:
                    st.success("Transcript fetched successfully.")

                    selected_output = output_map[output_type]

                    with st.spinner(f"Generating {selected_output}..."):
                        # Process with Gemini directly with fewer chunks
                        result = process_with_gemini(transcript, detail_level, selected_output, st.session_state.model_name)

                        # Store in session state based on type
                        if selected_output == "summary":
                            st.session_state.current_summary = result
                        elif selected_output == "notes":
                            st.session_state.current_notes = result
                        elif selected_output == "quiz":
                            st.session_state.current_quiz = result

                    # Add to history
                    add_to_history(url, video_title, selected_output)

                    # Display the result
                    if selected_output == "quiz":
                        display_quiz(result)
                    else:
                        display_content(result, selected_output, video_title)
                else:
                    st.error("Could not fetch transcript. Ensure the video has English captions.")
        else:
            st.error("Invalid YouTube URL. Please check and try again.")

def display_content(content, content_type, title):
    """Display content with appropriate formatting and download options"""
    type_emoji = {"summary": "📝", "notes": "📓"}
    type_name = {"summary": "Summary", "notes": "Study Notes"}

    st.markdown(f"### {type_emoji.get(content_type, '📄')} {type_name.get(content_type, 'Content')} - {title}")

    # Create tabs for different views
    tab1, tab2 = st.tabs(["Content", "Export Options"])

    with tab1:
        st.markdown(content)

    with tab2:
        col1, col2 = st.columns(2)
        with col1:
            # Fix for the f-string with regex issue
            safe_title = re.sub(r'[^\w]', '_', title)
            file_name = f"{content_type}_{safe_title}.docx"

            doc_file = save_to_word(content, file_name, title)
            st.download_button(
                "📥 Download as Word",
                open(doc_file, "rb"),
                file_name=file_name
            )

        with col2:
            # Copy to clipboard (simulated)
            st.button("📋 Copy to Clipboard",
                help="Click to copy content (simulated in this demo)")

def parse_quiz_content(quiz_content):
    """Parse quiz content and return a structured list of questions with options and answers"""
    questions = []

    # Use regex to extract different question types and their components
    # Extract multiple choice questions
    mc_section = re.search(r'# MULTIPLE CHOICE QUESTIONS(.*?)(?:# TRUE|$)', quiz_content, re.DOTALL)
    if mc_section:
        mc_questions = re.findall(r'(\d+\.\s+.*?)(?=\d+\.\s+|# TRUE|$)', mc_section.group(1), re.DOTALL)
        for q in mc_questions:
            question_text = re.search(r'\d+\.\s+(.*?)(?:\n[A-D]\)|\n$)', q, re.DOTALL)
            if question_text:
                question_dict = {
                    "type": "multiple_choice",
                    "text": question_text.group(1).strip(),
                    "options": []
                }

                # Extract options
                options = re.findall(r'([A-D]\))\s+(.*?)(?=\n[A-D]\)|\nCorrect:|\n$)', q, re.DOTALL)
                for option_letter, option_text in options:
                    question_dict["options"].append(f"{option_letter} {option_text.strip()}")

                # Extract correct answer
                correct = re.search(r'Correct:\s+([A-D])', q)
                if correct:
                    question_dict["correct"] = correct.group(1)

                questions.append(question_dict)

    # Extract true/false questions
    tf_section = re.search(r'# TRUE/FALSE QUESTIONS(.*?)(?:# SHORT|$)', quiz_content, re.DOTALL)
    if tf_section:
        tf_questions = re.findall(r'(\d+\.\s+.*?)(?=\d+\.\s+|# SHORT|$)', tf_section.group(1), re.DOTALL)
        for q in tf_questions:
            question_text = re.search(r'\d+\.\s+(.*?)(?:\nCorrect:|\n$)', q, re.DOTALL)
            if question_text:
                question_dict = {
                    "type": "true_false",
                    "text": question_text.group(1).strip()
                }

                # Extract correct answer
                correct = re.search(r'Correct:\s+(True|False)', q)
                if correct:
                    question_dict["correct"] = correct.group(1)

                questions.append(question_dict)

    # Extract short answer questions
    sa_section = re.search(r'# SHORT ANSWER QUESTIONS(.*?)(?=$)', quiz_content, re.DOTALL)
    if sa_section:
        sa_questions = re.findall(r'(\d+\.\s+.*?)(?=\d+\.\s+|$)', sa_section.group(1), re.DOTALL)
        for q in sa_questions:
            question_text = re.search(r'\d+\.\s+(.*?)(?:\nAnswer:|\n$)', q, re.DOTALL)
            if question_text:
                question_dict = {
                    "type": "short_answer",
                    "text": question_text.group(1).strip()
                }

                # Extract model answer
                answer = re.search(r'Answer:\s+(.*?)$', q, re.DOTALL)
                if answer:
                    question_dict["answer"] = answer.group(1).strip()

                questions.append(question_dict)

    return questions

def display_quiz(quiz_content):
    """Display quiz with interactive elements"""
    st.markdown(f"### 🧠 Quiz - {st.session_state.video_title}")

    # Create tabs for different views
    tab1, tab2 = st.tabs(["Take Quiz", "Export Options"])

    with tab1:
        # Parse quiz content and create input fields
        questions = parse_quiz_content(quiz_content)

        if not questions:
            st.error("Failed to parse quiz content properly. Displaying raw content instead.")
            st.markdown(quiz_content)
            return

        # Clear previous answers if needed
        if "user_answers" not in st.session_state or st.session_state.quiz_submitted:
            st.session_state.user_answers = {}
            st.session_state.quiz_submitted = False

        # Display questions in cards
        for i, question in enumerate(questions):
            with st.container():
                st.markdown(f"**Question {i+1}:** {question['text']}")

                if question['type'] == "multiple_choice":
                    options = [opt.strip() for opt in question['options']]
                    st.session_state.user_answers[i] = st.radio(
                        f"Select your answer for Question {i+1}:",
                        options,
                        key=f"q{i+1}",
                        label_visibility="collapsed"
                    )

                elif question['type'] == "true_false":
                    st.session_state.user_answers[i] = st.radio(
                        f"True or False for Question {i+1}:",
                        ["True", "False"],
                        key=f"q{i+1}",
                        label_visibility="collapsed"
                    )

                elif question['type'] == "short_answer":
                    st.session_state.user_answers[i] = st.text_area(
                        f"Your answer for Question {i+1}:",
                        key=f"q{i+1}",
                        height=100,
                        label_visibility="collapsed"
                    )

                st.markdown("---")

        if st.button("Submit Quiz", use_container_width=True):
            st.session_state.quiz_submitted = True
            display_quiz_results(questions)

    with tab2:
        col1, col2 = st.columns(2)
        with col1:
            # Fix for the f-string with regex issue
            safe_title = re.sub(r'[^\w]', '_', st.session_state.video_title)
            file_name = f"quiz_{safe_title}.docx"

            doc_file = save_to_word(quiz_content, file_name, f"Quiz - {st.session_state.video_title}")
            st.download_button(
                "📥 Download as Word",
                open(doc_file, "rb"),
                file_name=file_name
            )

        with col2:
            # Copy to clipboard (simulated)
            st.button("📋 Copy to Clipboard",
                help="Click to copy content (simulated in this demo)")

def display_quiz(quiz_content):
    """Display quiz with interactive elements"""
    st.markdown(f"### 🧠 Quiz - {st.session_state.video_title}")

    # Create tabs for different views
    tab1, tab2 = st.tabs(["Take Quiz", "Export Options"])

    with tab1:
        # Parse quiz content and create input fields
        questions = parse_quiz_content(quiz_content)

        if not questions:
            st.error("Failed to parse quiz content properly. Displaying raw content instead.")
            st.markdown(quiz_content)
            return

        # Clear previous answers if needed
        if "user_answers" not in st.session_state or st.session_state.quiz_submitted:
            st.session_state.user_answers = {}
            st.session_state.quiz_submitted = False

        # Display questions in cards
        for i, question in enumerate(questions):
            with st.container():
                st.markdown(f"**Question {i+1}:** {question['text']}")

                if question['type'] == "multiple_choice":
                    options = [opt.strip() for opt in question['options']]
                    st.session_state.user_answers[i] = st.radio(
                        f"Select your answer for Question {i+1}:",
                        options,
                        key=f"q{i+1}",
                        label_visibility="collapsed"
                    )

                elif question['type'] == "true_false":
                    st.session_state.user_answers[i] = st.radio(
                        f"True or False for Question {i+1}:",
                        ["True", "False"],
                        key=f"q{i+1}",
                        label_visibility="collapsed"
                    )

                elif question['type'] == "short_answer":
                    st.session_state.user_answers[i] = st.text_area(
                        f"Your answer for Question {i+1}:",
                        key=f"q{i+1}",
                        height=100,
                        label_visibility="collapsed"
                    )

                st.markdown("---")

        # Debug: Check if the button is clicked
        if st.button("Submit Quiz", use_container_width=True):
            st.session_state.quiz_submitted = True
            st.write("Quiz submitted. Displaying results...")  # Debug statement
            display_quiz_results(questions)

    with tab2:
        col1, col2 = st.columns(2)
        with col1:
            # Fix for the f-string with regex issue
            safe_title = re.sub(r'[^\w]', '_', st.session_state.video_title)
            file_name = f"quiz_{safe_title}.docx"

            doc_file = save_to_word(quiz_content, file_name, f"Quiz - {st.session_state.video_title}")
            st.download_button(
                "📥 Download as Word",
                open(doc_file, "rb"),
                file_name=file_name
            )

        with col2:
            # Copy to clipboard (simulated)
            st.button("📋 Copy to Clipboard",
                help="Click to copy content (simulated in this demo)")

def display_quiz_results(questions):
    """Display quiz results after submission"""
    st.markdown("### Quiz Results")

    correct_count = 0
    total_questions = len(questions)

    for i, question in enumerate(questions):
        user_answer = st.session_state.user_answers.get(i, "")

        with st.container():
            st.markdown(f"**Question {i+1}:** {question['text']}")
            st.markdown(f"*Your answer:* {user_answer}")

            if question['type'] in ["multiple_choice", "true_false"]:
                correct_answer = question.get('correct', '')
                # For multiple choice, we need to extract just the letter
                if question['type'] == "multiple_choice":
                    user_letter = user_answer.split(')')[0] if user_answer and ')' in user_answer else ""
                    is_correct = user_letter.strip() == correct_answer
                else:
                    is_correct = user_answer == correct_answer

                if is_correct:
                    st.markdown(f"<div class='correct-answer'>✓ Correct!</div>", unsafe_allow_html=True)
                    correct_count += 1
                else:
                    if question['type'] == "multiple_choice":
                        correct_option = next((opt for opt in question['options'] if opt.startswith(f"{correct_answer})")), "")
                        st.markdown(f"<div class='incorrect-answer'>✗ Incorrect. The correct answer is: {correct_option}</div>", unsafe_allow_html=True)
                    else:
                        st.markdown(f"<div class='incorrect-answer'>✗ Incorrect. The correct answer is: {correct_answer}</div>", unsafe_allow_html=True)
            else:
                # For short answer questions, show the model answer
                st.markdown(f"*Model answer:* {question.get('answer', 'No model answer available.')}")

            st.markdown("---")

    # Display final score
    score_percentage = (correct_count / total_questions) * 100 if total_questions > 0 else 0
    st.markdown(f"<div class='final-score'>Your Score: {correct_count}/{total_questions} ({score_percentage:.2f}%)</div>", unsafe_allow_html=True)

def render_history():
    """Render the history of processed videos"""
    st.markdown("### 🕰️ History")

    # Load history from file if not already in session state
    load_history()

    if not st.session_state.history:
        st.info("No history available. Process some videos to see your history here.")
    else:
        for entry in st.session_state.history:
            with st.container():
                st.markdown(f"**{entry['timestamp']}** - [{entry['title']}]({entry['url']})")
                st.markdown(f"*Type:* {entry['type'].capitalize()}")
                st.markdown("---")

def render_settings():
    """Render settings page"""
    st.markdown("### ⚙️ Settings")
    st.info("Settings page is under development. Stay tuned for more customization options!")

def main():
    """Main function to control the app flow"""
    selected = render_sidebar()

    if selected == "Summarizer":
        render_header()
        render_main_app()
    elif selected == "History":
        render_header()
        render_history()
    elif selected == "Settings":
        render_header()
        render_settings()

if __name__ == "__main__":
    main()